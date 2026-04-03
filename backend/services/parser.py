import io
from typing import Optional


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF bytes. Tries pdfplumber first, falls back to PyMuPDF."""
    # Try pdfplumber first
    try:
        import pdfplumber
        text_parts = []
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                t = page.extract_text()
                if t:
                    text_parts.append(t)
        result = "\n".join(text_parts)
        if result.strip():
            return result
    except Exception:
        pass

    # Fallback: PyMuPDF (fitz)
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        text_parts = []
        for page in doc:
            t = page.get_text()
            if t:
                text_parts.append(t)
        doc.close()
        result = "\n".join(text_parts)
        if result.strip():
            return result
    except Exception:
        pass

    raise ValueError("PDF parsing failed: could not extract text. File may be scanned/image-based or corrupted.")


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX bytes using python-docx."""
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        result = "\n".join(paragraphs)
        if result.strip():
            return result
        raise ValueError("DOCX appears to be empty.")
    except ValueError:
        raise
    except Exception as e:
        raise ValueError(f"DOCX parsing failed: {e}")


def extract_text_from_doc(file_bytes: bytes) -> str:
    """Extract text from legacy .doc files using antiword or textract fallback."""
    # Try python-docx anyway (sometimes works with .doc)
    try:
        return extract_text_from_docx(file_bytes)
    except Exception:
        pass

    # Try reading as plain text (some .doc files are RTF or plain)
    try:
        text = file_bytes.decode("utf-8", errors="ignore")
        # Strip RTF control sequences if present
        import re
        if text.startswith("{\\rtf"):
            # Basic RTF strip
            text = re.sub(r'\{[^{}]*\}', ' ', text)
            text = re.sub(r'\\[a-z]+\d* ?', ' ', text)
            text = re.sub(r'[{}\\]', ' ', text)
        clean = " ".join(text.split())
        if len(clean) > 50:  # Minimum viable content
            return clean
    except Exception:
        pass

    raise ValueError(".doc file could not be parsed. Please save as .docx or .pdf and re-upload.")


def extract_text_from_txt(file_bytes: bytes) -> str:
    """Extract text from plain text file."""
    for encoding in ("utf-8", "latin-1", "cp1252"):
        try:
            return file_bytes.decode(encoding)
        except Exception:
            continue
    raise ValueError("Could not decode text file.")


def extract_text(file_bytes: bytes, filename: str) -> str:
    """
    Auto-detect file type and extract text.
    Supported: PDF, DOCX, DOC, TXT
    """
    fname_lower = filename.lower()

    if fname_lower.endswith(".pdf"):
        return extract_text_from_pdf(file_bytes)
    elif fname_lower.endswith(".docx"):
        return extract_text_from_docx(file_bytes)
    elif fname_lower.endswith(".doc"):
        return extract_text_from_doc(file_bytes)
    elif fname_lower.endswith(".txt"):
        return extract_text_from_txt(file_bytes)
    else:
        # Last resort: try each parser
        errors = []
        for fn, label in [
            (extract_text_from_pdf, "PDF"),
            (extract_text_from_docx, "DOCX"),
            (extract_text_from_txt, "TXT"),
        ]:
            try:
                result = fn(file_bytes)
                if result.strip():
                    return result
            except Exception as e:
                errors.append(f"{label}: {e}")
        raise ValueError(
            f"Unsupported or unreadable file: {filename}. "
            f"Supported formats: PDF, DOCX, DOC, TXT. Errors: {'; '.join(errors)}"
        )
