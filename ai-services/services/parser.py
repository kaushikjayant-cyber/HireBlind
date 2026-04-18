"""
parser.py — Resume text extraction with multi-library PDF fallback chain.

PDF extraction priority:
  1. pdfplumber  (best layout awareness, tables)
  2. PyMuPDF / fitz  (fast, great Unicode support)
  3. pdfminer.six  (pure-Python fallback)
  4. Raw byte scan  (emergency last resort for text PDFs)

DOCX: python-docx (paragraphs + tables)
TXT:  UTF-8 → latin-1 → cp1252 → ascii(replace)
"""
import io
import re
import logging

logger = logging.getLogger(__name__)


# ── Public entry point ────────────────────────────────────────────────────────

def extract_text(file_bytes: bytes, filename: str) -> str:
    """
    Extract raw text from a resume file.

    Args:
        file_bytes: Raw bytes of the uploaded file.
        filename:   Original filename (used to determine type).

    Returns:
        Extracted text string.

    Raises:
        ValueError: Unsupported / corrupt / empty file.
    """
    if not file_bytes:
        raise ValueError("File is empty.")

    fname = filename.lower().strip()

    if fname.endswith(".pdf"):
        return _extract_pdf(file_bytes, filename)
    elif fname.endswith(".docx"):
        return _extract_docx(file_bytes, filename)
    elif fname.endswith(".doc"):
        raise ValueError(
            f"'{filename}' is in legacy .doc format. "
            "Please re-save as .docx or .pdf and re-upload."
        )
    elif fname.endswith(".txt"):
        return _extract_txt(file_bytes)
    else:
        ext = fname.rsplit(".", 1)[-1] if "." in fname else "unknown"
        raise ValueError(
            f"Unsupported file type '.{ext}'. Accepted: PDF, DOCX, TXT."
        )


# ── PDF extraction with fallback chain ───────────────────────────────────────

def _extract_pdf(file_bytes: bytes, filename: str) -> str:
    """Try multiple PDF libraries in order. Raise only if all fail."""
    errors = []

    # ── Attempt 1: pdfplumber ────────────────────────────────────────────────
    try:
        import pdfplumber
        logger.info(f"[parser] PDF: trying pdfplumber for '{filename}'")
        parts = []
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    parts.append(text.strip())
                for table in (page.extract_tables() or []):
                    for row in table:
                        row_text = " | ".join(
                            cell.strip() for cell in row if cell and cell.strip()
                        )
                        if row_text:
                            parts.append(row_text)
        result = "\n".join(parts)
        if result.strip():
            logger.info(f"[parser] pdfplumber extracted {len(result)} chars")
            return result
        logger.warning("[parser] pdfplumber returned empty text, trying next")
    except ImportError:
        logger.warning("[parser] pdfplumber not installed — trying PyMuPDF")
    except Exception as e:
        errors.append(f"pdfplumber: {e}")
        logger.warning(f"[parser] pdfplumber failed: {e}")

    # ── Attempt 2: PyMuPDF (fitz) ────────────────────────────────────────────
    try:
        import fitz  # PyMuPDF
        logger.info(f"[parser] PDF: trying PyMuPDF for '{filename}'")
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        parts = []
        for page in doc:
            text = page.get_text("text")
            if text and text.strip():
                parts.append(text.strip())
        doc.close()
        result = "\n".join(parts)
        if result.strip():
            logger.info(f"[parser] PyMuPDF extracted {len(result)} chars")
            return result
        logger.warning("[parser] PyMuPDF returned empty text, trying next")
    except ImportError:
        logger.warning("[parser] PyMuPDF not installed — trying pdfminer")
    except Exception as e:
        errors.append(f"PyMuPDF: {e}")
        logger.warning(f"[parser] PyMuPDF failed: {e}")

    # ── Attempt 3: pdfminer.six ──────────────────────────────────────────────
    try:
        from pdfminer.high_level import extract_text as pdfminer_extract
        logger.info(f"[parser] PDF: trying pdfminer for '{filename}'")
        result = pdfminer_extract(io.BytesIO(file_bytes))
        if result and result.strip():
            logger.info(f"[parser] pdfminer extracted {len(result)} chars")
            return result.strip()
        logger.warning("[parser] pdfminer returned empty text, trying next")
    except ImportError:
        logger.warning("[parser] pdfminer.six not installed — trying raw scan")
    except Exception as e:
        errors.append(f"pdfminer: {e}")
        logger.warning(f"[parser] pdfminer failed: {e}")

    # ── Attempt 4: Raw byte text scan (emergency fallback) ───────────────────
    try:
        logger.info(f"[parser] PDF: trying raw byte scan for '{filename}'")
        raw = _raw_pdf_text_scan(file_bytes)
        if raw.strip():
            logger.info(f"[parser] Raw scan extracted {len(raw)} chars")
            return raw
    except Exception as e:
        errors.append(f"raw_scan: {e}")

    # All methods failed
    error_summary = "; ".join(errors) if errors else "No text layer found"
    raise RuntimeError(
        f"PDF parsing failed for '{filename}'. {error_summary}. "
        "If this is a scanned PDF, please use a text-based PDF or DOCX."
    )


def _raw_pdf_text_scan(file_bytes: bytes) -> str:
    """
    Emergency fallback: scan raw PDF bytes for BT...ET text blocks.
    Works only for simple unencrypted PDFs with embedded text streams.
    """
    text = file_bytes.decode("latin-1", errors="replace")
    # Extract text between BT (Begin Text) and ET (End Text) markers
    blocks = re.findall(r"BT\s*(.*?)\s*ET", text, re.DOTALL)
    words = []
    for block in blocks:
        # Tj operator: (text) Tj
        tj_matches = re.findall(r"\((.*?)\)\s*Tj", block)
        words.extend(tj_matches)
        # TJ operator: [(text) ...] TJ
        tj_array = re.findall(r"\((.*?)\)", block)
        words.extend(tj_array)

    cleaned = []
    for w in words:
        w = w.replace("\\n", "\n").replace("\\r", "").replace("\\t", " ")
        if w.strip():
            cleaned.append(w)
    return " ".join(cleaned)


# ── DOCX extraction ───────────────────────────────────────────────────────────

def _extract_docx(file_bytes: bytes, filename: str) -> str:
    try:
        from docx import Document
    except ImportError:
        raise RuntimeError(
            "python-docx is not installed. Run: pip install python-docx"
        )

    try:
        doc = Document(io.BytesIO(file_bytes))
    except Exception as e:
        raise ValueError(
            f"Could not open '{filename}' as DOCX. "
            f"File may be corrupted. Error: {e}"
        )

    parts = []
    for para in doc.paragraphs:
        t = para.text.strip()
        if t:
            parts.append(t)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                t = cell.text.strip()
                if t and t not in parts:
                    parts.append(t)

    return "\n".join(parts)


# ── TXT extraction ────────────────────────────────────────────────────────────

def _extract_txt(file_bytes: bytes) -> str:
    for encoding in ("utf-8", "utf-8-sig", "latin-1", "cp1252"):
        try:
            return file_bytes.decode(encoding)
        except (UnicodeDecodeError, LookupError):
            continue
    return file_bytes.decode("ascii", errors="replace")
